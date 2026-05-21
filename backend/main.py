<<<<<<< HEAD
from fastapi import FastAPI, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
=======
# This connects everything

from fastapi import FastAPI, UploadFile, Form
>>>>>>> 210088843dd3d6d075a62a811f9a497aa0aa3155
import pdfplumber
from docx import Document

app = FastAPI()

<<<<<<< HEAD
# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Extract Resume Text
def extract_text(file: UploadFile):

=======
def extract_text(file: UploadFile):
>>>>>>> 210088843dd3d6d075a62a811f9a497aa0aa3155
    filename = file.filename.lower()
    text = ""

    try:
<<<<<<< HEAD

        if filename.endswith(".pdf"):

            with pdfplumber.open(file.file) as pdf:

                for page in pdf.pages:

                    page_text = page.extract_text()

=======
        if filename.endswith(".pdf"):
            with pdfplumber.open(file.file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
>>>>>>> 210088843dd3d6d075a62a811f9a497aa0aa3155
                    if page_text:
                        text += page_text

        elif filename.endswith(".docx"):
<<<<<<< HEAD

            doc = Document(file.file)

=======
            doc = Document(file.file)
>>>>>>> 210088843dd3d6d075a62a811f9a497aa0aa3155
            text = "\n".join([p.text for p in doc.paragraphs])

        return text.strip()

    except:
        return ""
<<<<<<< HEAD

# Analyze API
@app.post("/analyze")
def analyze(resume: UploadFile, job_desc: str = Form(...)):

    # File Validation
    if not (
        resume.filename.endswith(".pdf")
        or resume.filename.endswith(".docx")
    ):
        return {
            "error": "Only PDF or DOCX files allowed"
        }

    text = extract_text(resume)

    if not text:
        return {
            "error": "Could not read resume"
        }

    # Dummy ATS Logic
=======
    

@app.post("/analyze")
def analyze(resume: UploadFile, job_desc: str = Form(...)):

    text = extract_text(resume)

    # SAFE CHECK (THIS PREVENTS 500 ERROR)
    if not text:
        return {
            "error": "Could not read resume file. Try another PDF/DOCX."
        }

    # SIMPLE DUMMY SCORE (so project ALWAYS runs)
>>>>>>> 210088843dd3d6d075a62a811f9a497aa0aa3155
    score = 70

    if "python" in job_desc.lower():
        score += 10
<<<<<<< HEAD

    if "fastapi" in job_desc.lower():
        score += 10

    if "sql" in job_desc.lower():
        score += 5

    score = min(score, 100)

    # Match Level
    if score >= 80:
        level = "Strong Match"
    elif score >= 50:
        level = "Medium Match"
    else:
        level = "Weak Match"

    return {
        "ats_score": score,
        "match_level": level
    }
=======
    if "fastapi" in job_desc.lower():
        score += 10

    return {
        "ats_score": min(score, 100),
        "match_level": "Medium Match"
    }
   
from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
>>>>>>> 210088843dd3d6d075a62a811f9a497aa0aa3155
